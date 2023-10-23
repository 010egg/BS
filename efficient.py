
import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import urljoin, urlparse


def get_all_links(url):
    try:
        # 发送HTTP请求
        page = requests.get(url)

        # 解析页面
        soup = BeautifulSoup(page.content, 'html.parser')

        # 查找页面中的所有链接
        links = soup.find_all('a', href=True)

        # 筛选出有效的链接并获取完整的链接地址
        urls = []
        for link in links:
            href = link['href']
            parsed_href = urlparse(href)
            # 排除JavaScript操作、空链接等非HTTP链接
            if parsed_href.scheme in ['http', 'https']:
                urls.append(urljoin(url, href))

        return urls
    except Exception as e:
        print(f"An error occurred while fetching links: {e}")
        return []
# 目标网站URL
def fetch_title_and_body(url):
    try:
        # 获取单个页面的内容
        page = requests.get(url)

        # 解析HTML
        soup = BeautifulSoup(page.content, 'html.parser')

        # 获取标题
        title = soup.find('title').get_text()

        # 如果标题中不包含 "Spark -"，则不处理该页面的内容
        if "Spark -" not in title:
            return None, None  # 不是我们要找的页面

        # 获取body内容（这可能稍微复杂一些，因为它取决于网站的结构）
        body = soup.find('body').get_text()

        return title, body

    except Exception as e:
        print(f"An error occurred: {e}")
        return None, None

def write_to_word(document, title, body):

    # 将标题和内容添加到Word文档
    document.add_heading(title, level=1)
    document.add_paragraph(body)
    # 主程序开始运行的地方
if __name__ == "__main__":
    url = 'https://study.sf.163.com/documents/read/service_support/dsc-t-03'
    urls = get_all_links(url)
    doc = Document()
    for link in urls:
        title, body = fetch_title_and_body(urls)
        if title and body:  # 如果标题和正文内容存在
            print(f"Page Title: {title}\n")
            print(f"Body Content: \n{body}")