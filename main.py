
import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import urljoin

def get_all_links(url):
        # 发送一个HTTP请求
        page = requests.get(url)

        # 使用BeautifulSoup解析页面
        soup = BeautifulSoup(page.content, 'html.parser')

        # 查找页面中的所有链接
        links = soup.find_all('a', href=True)

        # 获取完整的链接地址
        urls = [urljoin(url, link['href']) for link in links]

        return urls
# 目标网站URL
def fetch_title_and_body(url):
    try:
        # 获取单个页面的内容
        page = requests.get(url)

        # 解析HTML
        soup = BeautifulSoup(page.content, 'html.parser')

        # 获取标题
        title = soup.find('title').get_text()

        # 如果标题中不包含 "Spark -"，则不处理该页面的内容
        if "Spark -" not in title:
            return None, None  # 不是我们要找的页面

        # 获取body内容（这可能稍微复杂一些，因为它取决于网站的结构）
        body = soup.find('body').get_text()

        return title, body

    except Exception as e:
        print(f"An error occurred: {e}")
        return None, None

def write_to_word(document, title, body):

    # 将标题和内容添加到Word文档
    document.add_heading(title, level=1)
    document.add_paragraph(body)
    # 主程序开始运行的地方
# if __name__ == "__main__":
#     url = 'https://study.sf.163.com/documents/read/service_support/dsc-t-03'
#     urls = get_all_links(url)
#
#     for link in urls:
#         title, body = fetch_title_and_body(link)
#         if title and body:  # 如果标题和正文内容存在
#             print(f"Page Title: {title}")


def is_valid_link(url, base_url):
    # 只处理base_url开头的链接，避免外部链接
    return url.startswith(base_url)


if __name__ == "__main__":
    base_url = 'https://study.sf.163.com'
    start_url = 'https://study.sf.163.com/documents/read/service_support/dsc-t-03'

    urls = get_all_links(start_url)

    # 创建一个Word文档对象
    doc = Document()

    for link in urls:
        if is_valid_link(link, base_url):
            title, body = fetch_title_and_body(link)
            if title and body:  # 如果标题和正文内容存在
                print(f"Page Title: {title}")
                write_to_word(doc, title, body)  # 将内容写入Word文档

    # 保存Word文档
    doc.save("spark_content.docx")



